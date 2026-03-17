from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, BackgroundTasks, Form
from sqlalchemy.orm import Session, joinedload
import os
import json
from datetime import datetime, timezone
from app.infrastructure.database import get_db, SessionLocal
from app.domain.models import User, Application, Job, ResumeExtraction, Interview, InterviewAnswer
from app.domain.schemas import ApplicationCreate, ApplicationStatusUpdate, ApplicationResponse, ApplicationDetailResponse, ApplicationNotesUpdate
from app.core.auth import get_current_user, get_current_hr
from app.services.ai_service import parse_resume_with_ai
from app.services.email_service import send_application_received_email, send_rejected_email, send_approved_for_interview_email
import secrets
from passlib.context import CryptContext
from datetime import datetime, timedelta

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

from app.core.config import get_settings
settings = get_settings()

UPLOAD_DIR = settings.uploads_dir / "resumes"
PHOTO_DIR = settings.uploads_dir / "photos"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
PHOTO_DIR.mkdir(parents=True, exist_ok=True)

router = APIRouter(prefix="/api/applications", tags=["applications"])

@router.get("/ranking/{job_id}")
def get_candidate_ranking(
    job_id: int,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Get ranked candidates for a specific job (Point 3)"""
    from app.services.candidate_service import CandidateService
    service = CandidateService(db)
    ranked = service.get_ranked_candidates(job_id)
    
    result = []
    for idx, app in enumerate(ranked):
        result.append({
            "rank": idx + 1,
            "id": app.id,
            "candidate_name": app.candidate_name,
            "composite_score": app.composite_score,
            "recommendation": app.recommendation,
            "status": app.status
        })
    return result


@router.post("/apply", response_model=ApplicationResponse)
async def apply_for_job(
    job_id: int = Form(...),
    candidate_name: str = Form(...),
    candidate_email: str = Form(...),
    candidate_phone: str = Form(None),
    resume_file: UploadFile = File(...),
    photo_file: UploadFile = File(None),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: Session = Depends(get_db)
):
    """Apply for a job with resume (Public endpoint)"""
    # Check if job exists and is open
    job = db.query(Job).filter(Job.id == job_id, Job.status == "open").first()
    if not job:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Job not found or not open"
        )
    
    # Check if already applied
    # Check if already applied
    candidate_email = candidate_email.lower().strip()
    existing_app = db.query(Application).filter(
        Application.job_id == job_id,
        Application.candidate_email == candidate_email
    ).first()
    
    if existing_app:
        # If the previous application was rejected, allow re-application by deleting the old one
        if existing_app.status == "rejected":
            try:
                # Start fresh - delete the old application tree (cascades should handle relations)
                db.delete(existing_app)
                db.commit()
            except Exception as e:
                db.rollback()
                raise HTTPException(status_code=500, detail="Failed to recycle rejected application securely")
            # Loop continues to create new app
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="You have already applied for this job"
            )
    
    MAX_FILE_SIZE = 5 * 1024 * 1024 # 5MB
    ALLOWED_TYPES = ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "text/plain"]
    
    # Validate file content type
    if resume_file.content_type not in ALLOWED_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid file type. Only PDF and DOCX allowed."
        )
        
    # Validate file size (Need to read chunk to be safe, but spooled file has .size or we check after reading)
    # Since UploadFile is spooled, we can check size if headers provided, or read content.
    content = await resume_file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File too large. Maximum size is 5MB."
        )
            
    # Save resume file
    file_extension = resume_file.filename.split(".")[-1]
    safe_email = candidate_email.replace('@', '_').replace('.', '_')
    filename = f"{safe_email}_{job_id}_{datetime.now(timezone.utc).timestamp()}.{file_extension}"
    
    # Absolute path for saving the file
    abs_file_path = os.path.join(UPLOAD_DIR, filename).replace("\\", "/")
    # Relative path for storing in DB (starts with 'uploads/')
    rel_file_path = f"uploads/resumes/{filename}"
    
    with open(abs_file_path, "wb") as f:
        f.write(content)
    
    # Save photo file if provided
    rel_photo_path = None
    if photo_file:
        photo_content = await photo_file.read()
        if len(photo_content) > MAX_FILE_SIZE:
             raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Photo too large. Maximum size is 5MB."
            )
        
        photo_ext = photo_file.filename.split(".")[-1]
        photo_filename = f"photo_{safe_email}_{job_id}_{datetime.now(timezone.utc).timestamp()}.{photo_ext}"
        
        # Absolute path for saving
        abs_photo_path = os.path.join(PHOTO_DIR, photo_filename).replace("\\", "/")
        # Relative path for DB
        rel_photo_path = f"uploads/photos/{photo_filename}"
        
        with open(abs_photo_path, "wb") as f:
            f.write(photo_content)
    
    # Create application
    new_application = Application(
        job_id=job_id,
        candidate_name=candidate_name,
        candidate_email=candidate_email,
        candidate_phone=candidate_phone,
        resume_file_path=rel_file_path,
        resume_file_name=resume_file.filename,
        candidate_photo_path=rel_photo_path,
        status="applied"
    )
    
    try:
        db.add(new_application)
        db.commit()
        db.refresh(new_application)
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail="Failed to save new application securely")
    
    # 1. Notify candidate immediately of receipt
    background_tasks.add_task(send_application_received_email, candidate_email, job.title)

    # 2. Move all heavy processing to background task to prevent timeouts (Point 10 - Robustness)
    background_tasks.add_task(
        process_application_background, 
        new_application.id, 
        job_id, 
        abs_file_path, 
        candidate_email, 
        candidate_name
    )
    
    return new_application

async def process_application_background(application_id: int, job_id: int, abs_file_path: str, candidate_email: str, candidate_name: str):
    """Heavy AI processing and notification workflow in background"""
    db = SessionLocal()
    try:
        from app.services.candidate_service import CandidateService
        cand_service = CandidateService(db)
        
        # Reload objects in this session
        application = db.query(Application).filter(Application.id == application_id).first()
        job = db.query(Job).filter(Job.id == job_id).first()
        if not application or not job:
            db.close()
            return

        # 1. Initial State
        cand_service.advance_stage(application_id, "Application Submitted", "pass")
        cand_service.create_audit_log(None, "APPLICATION_SUBMITTED", "Application", application_id, {"email": candidate_email})
        
        # 2. Screening Stage
        cand_service.advance_stage(application_id, "Resume Screening", "pending")
        
        # Parse resume text based on file type
        resume_text = ""
        try:
            file_ext = abs_file_path.lower().split('.')[-1]
            if file_ext == 'pdf':
                from pypdf import PdfReader
                if not os.path.exists(abs_file_path):
                    print(f"[ERROR] Resume file not found at: {abs_file_path}")
                    raise FileNotFoundError(f"Resume file not found: {abs_file_path}")
                reader = PdfReader(abs_file_path)
                for page in reader.pages:
                    resume_text += page.extract_text() + "\n"
            elif file_ext in ['docx', 'doc']:
                import docx
                doc = docx.Document(abs_file_path)
                for para in doc.paragraphs:
                    resume_text += para.text + "\n"
            else:
                with open(abs_file_path, "rb") as f:
                    resume_text = f.read().decode('utf-8', errors='ignore')
        except Exception as e:
            print(f"Background Text Extraction Error: {e}")
            cand_service.create_audit_log(None, "RESUME_TEXT_EXTRACTION_FAILED", "Application", application_id, {"error": str(e)})
            resume_text = "Error extracting text."
        
        if not resume_text.strip():
            resume_text = "No readable text found."

        # AI Parsing
        extraction_data = await parse_resume_with_ai(resume_text, job_id, job.description)
        
        # Store extraction
        resume_extraction = ResumeExtraction(
            application_id=application_id,
            extracted_text=resume_text,
            summary=extraction_data.get("summary", ""),
            extracted_skills=json.dumps(extraction_data.get("skills") or []),
            years_of_experience=extraction_data.get("experience"),
            education=json.dumps(extraction_data.get("education") or []),
            previous_roles=json.dumps(extraction_data.get("roles") or []),
            experience_level=extraction_data.get("experience_level"),
            resume_score=extraction_data.get("score", 0),
            skill_match_percentage=extraction_data.get("match_percentage", 0)
        )
        db.add(resume_extraction)
        
        # Update Application summary fields
        application.resume_score = extraction_data.get("score", 0)
        db.commit()

        # ── NO AUTOMATIC DECISION ──
        # All decisions (approve, reject) are now manual HR actions only.
        # The application stays in 'applied' state until HR takes action.
        cand_service.advance_stage(application_id, "Resume Screening", "pass", 
                                   extraction_data.get("score", 0) * 10, 
                                   "AI analysis complete — awaiting HR decision")
        cand_service.create_audit_log(None, "RESUME_SCREENING_COMPLETED", "Application", application_id, 
                                      {"score": extraction_data.get("score", 0), "match": extraction_data.get("match_percentage", 0)})
        
        db.commit()
    except Exception as e:
        print(f"CRITICAL Background Error processing application {application_id}: {e}")
        db.rollback()
        # Log the critical error
        try:
            cand_service = CandidateService(db) # Re-initialize if needed, or pass db
            cand_service.create_audit_log(None, "BACKGROUND_PROCESSING_FAILED", "Application", application_id, {"error": str(e)})
            # Optionally update application status to indicate processing failed
            application = db.query(Application).filter(Application.id == application_id).first()
            if application:
                application.status = "applied"  # Keep in 'applied' — HR can retry or proceed manually
                application.hr_notes = f"AI analysis failed: {e}"
                db.commit()
        except Exception as log_e:
            print(f"Failed to log critical error for application {application_id}: {log_e}")
    finally:
        db.close()
    
    # The return value of a background task is not used by FastAPI.
    # The original snippet returned new_application, but it's not necessary here.
    # Keeping it for consistency with the provided snippet, but it has no effect.
    return application 

@router.get("", response_model=list[ApplicationDetailResponse])
def get_hr_applications(
    job_id: int = None,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Get all applications for HR's jobs (HR only)"""
    # Join with stages for visualization (Point 12)
    # Use outerjoin to ensure apps with missing jobs (shouldn't happen but safe) still show or at least don't crash
    query = db.query(Application).outerjoin(Job).options(
        joinedload(Application.job),
        joinedload(Application.resume_extraction),
        joinedload(Application.interview),
        joinedload(Application.pipeline_stages)
    )
    
    # Filter by job if requested
    if job_id:
        query = query.filter(Application.job_id == job_id)
        
    # Security: Only admins can see everything. Others see their own jobs' apps.
    if current_user.role != "admin":
        print(f"DEBUG: Filtering applications for HR ID {current_user.id}")
        query = query.filter(Job.hr_id == current_user.id)
    else:
        print("DEBUG: Admin viewing all applications")
        
    applications = query.all()
    print(f"DEBUG: Found {len(applications)} applications for user {current_user.id}")
    
    for app in applications:
        if app.candidate_photo_path and ":" in app.candidate_photo_path:
            idx = app.candidate_photo_path.find("uploads")
            if idx != -1:
                app.candidate_photo_path = app.candidate_photo_path[idx:].replace("\\", "/")
        if app.resume_file_path and ":" in app.resume_file_path:
            idx = app.resume_file_path.find("uploads")
            if idx != -1:
                app.resume_file_path = app.resume_file_path[idx:].replace("\\", "/")
    
    return applications

@router.get("/{application_id}", response_model=ApplicationDetailResponse)
def get_application(
    application_id: int,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Get application details (HR only)"""
    application = db.query(Application).options(
        joinedload(Application.job),
        joinedload(Application.resume_extraction),
        joinedload(Application.interview),
        joinedload(Application.pipeline_stages)
    ).filter(Application.id == application_id).first()
    
    if not application:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Application not found"
        )
    
    # Sanitize paths
    if application.candidate_photo_path and ":" in application.candidate_photo_path:
        idx = application.candidate_photo_path.find("uploads")
        if idx != -1:
            application.candidate_photo_path = application.candidate_photo_path[idx:].replace("\\", "/")
    if application.resume_file_path and ":" in application.resume_file_path:
        idx = application.resume_file_path.find("uploads")
        if idx != -1:
            application.resume_file_path = application.resume_file_path[idx:].replace("\\", "/")
    
    return application

async def retry_application_background(application_id: int, job_id: int, abs_file_path: str):
    """Safely retry AI resume extraction without altering pipeline stages or triggering emails."""
    db = SessionLocal()
    try:
        from app.services.candidate_service import CandidateService
        cand_service = CandidateService(db)
        
        application = db.query(Application).filter(Application.id == application_id).first()
        job = db.query(Job).filter(Job.id == job_id).first()
        if not application or not job:
            db.close()
            return
            
        print(f"Retrying AI extraction for application {application_id}...")
        
        # Parse resume text based on file type
        resume_text = ""
        try:
            file_ext = abs_file_path.lower().split('.')[-1]
            if file_ext == 'pdf':
                from pypdf import PdfReader
                reader = PdfReader(abs_file_path)
                for page in reader.pages:
                    resume_text += page.extract_text() + "\n"
            elif file_ext in ['docx', 'doc']:
                import docx
                doc = docx.Document(abs_file_path)
                for para in doc.paragraphs:
                    resume_text += para.text + "\n"
            else:
                with open(abs_file_path, "rb") as f:
                    resume_text = f.read().decode('utf-8', errors='ignore')
        except Exception as e:
            print(f"Retry Text Extraction Error: {e}")
            cand_service.create_audit_log(None, "RETRY_TEXT_EXTRACTION_FAILED", "Application", application_id, {"error": str(e)})
            resume_text = "Error extracting text."
        
        if not resume_text.strip():
            resume_text = "No readable text found."

        # AI Parsing
        extraction_data = await parse_resume_with_ai(resume_text, job_id, job.description)
        
        # Look for existing extraction or create new
        resume_extraction = db.query(ResumeExtraction).filter(ResumeExtraction.application_id == application_id).first()
        if not resume_extraction:
            resume_extraction = ResumeExtraction(application_id=application_id)
            db.add(resume_extraction)
            
        resume_extraction.extracted_text = resume_text
        resume_extraction.summary = extraction_data.get("summary", "")
        resume_extraction.extracted_skills = json.dumps(extraction_data.get("skills") or [])
        resume_extraction.years_of_experience = extraction_data.get("experience")
        resume_extraction.education = json.dumps(extraction_data.get("education") or [])
        resume_extraction.previous_roles = json.dumps(extraction_data.get("roles") or [])
        resume_extraction.experience_level = extraction_data.get("experience_level")
        resume_extraction.resume_score = extraction_data.get("score", 0)
        resume_extraction.skill_match_percentage = extraction_data.get("match_percentage", 0)
        
        application.resume_score = extraction_data.get("score", 0)
        
        # Application should already be in 'applied' state, ensure it stays there
        # (processing_failed is removed from FSM)
             
        db.commit()
        cand_service.create_audit_log(None, "AI_ANALYSIS_RETRY_SUCCESS", "Application", application_id, {"score": extraction_data.get("score")})
        print(f"Retry successful for application {application_id}")
    except Exception as e:
        print(f"Retry Background Error processing application {application_id}: {e}")
        db.rollback()
        try:
            cand_service = CandidateService(db)
            cand_service.create_audit_log(None, "AI_ANALYSIS_RETRY_FAILED", "Application", application_id, {"error": str(e)})
        except:
             pass
    finally:
        db.close()

@router.post("/{application_id}/retry-analysis")
async def retry_resume_analysis(
    application_id: int, 
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Manually trigger AI resume analysis if it failed"""
    application = db.query(Application).filter(Application.id == application_id).first()
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
        
    abs_file_path = ""
    if application.resume_file_path:
        # Resolve the actual file path using application settings
        # In the DB, it's typically 'uploads/resumes/filename.pdf'
        # settings.uploads_dir is 'D:\CALDIM\ars\RIMS\backend\uploads'
        clean_path = application.resume_file_path
        if clean_path.startswith("uploads/"):
             clean_path = clean_path[8:]
        elif clean_path.startswith("uploads\\"):
             clean_path = clean_path[8:]
        
        abs_file_path = str(settings.uploads_dir / clean_path)
        
    if not abs_file_path or not os.path.exists(abs_file_path):
         print(f"Error locating file path: DB path='{application.resume_file_path}', Tried path='{abs_file_path}'")
         raise HTTPException(status_code=400, detail="Resume file not found on server")
         
    background_tasks.add_task(
        retry_application_background, 
        application.id, 
        application.job_id, 
        abs_file_path
    )
    
    return {"status": "success", "message": "Analysis restarted in background safely"}

@router.put("/{application_id}/status")


def update_application_status(
    application_id: int,
    status_update: ApplicationStatusUpdate,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """
    Execute a state transition on an application via the finite state machine.
    
    The frontend sends an 'action' (e.g. 'approve_for_interview', 'reject', 
    'call_for_interview', 'review_later', 'hire').  The FSM validates the 
    transition, updates the status atomically, logs the change, and returns 
    the result including which email to send.
    """
    from app.services.state_machine import (
        CandidateStateMachine, TransitionAction,
        InvalidTransitionError, DuplicateTransitionError,
    )

    application = db.query(Application).filter(Application.id == application_id).first()
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")

    # Parse the action
    try:
        action = TransitionAction(status_update.action)
    except ValueError:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid action: '{status_update.action}'. "
                   f"Valid actions: {[a.value for a in TransitionAction if not a.value.startswith('system_')]}"
        )

    # Block system actions from HR endpoint
    if action.value.startswith("system_"):
        raise HTTPException(status_code=400, detail="System actions cannot be triggered manually")

    # Execute FSM transition
    fsm = CandidateStateMachine(db)
    try:
        result = fsm.transition(
            application=application,
            action=action,
            user_id=current_user.id,
            notes=status_update.hr_notes,
        )
    except InvalidTransitionError as e:
        raise HTTPException(status_code=400, detail=e.message)
    except DuplicateTransitionError as e:
        raise HTTPException(status_code=400, detail=str(e))

    # Handle HR notes
    if status_update.hr_notes:
        application.hr_notes = status_update.hr_notes

    # ─── Post-transition side effects ───────────────────────────────────
    raw_access_key = None

    if action == TransitionAction.APPROVE_FOR_INTERVIEW:
        # Create or refresh interview record + access key
        raw_access_key = _ensure_interview_record(application, db)

    if action == TransitionAction.HIRE:
        # Create HiringDecision record
        from app.domain.models import HiringDecision
        existing = db.query(HiringDecision).filter(
            HiringDecision.application_id == application_id
        ).first()
        if not existing:
            decision = HiringDecision(
                application_id=application_id,
                hr_id=current_user.id,
                decision="hired",
                decision_comments=status_update.hr_notes or "Hired via pipeline",
                decided_at=datetime.now(timezone.utc),
            )
            db.add(decision)

    if action == TransitionAction.REJECT:
        from app.domain.models import HiringDecision
        existing = db.query(HiringDecision).filter(
            HiringDecision.application_id == application_id
        ).first()
        if not existing:
            decision = HiringDecision(
                application_id=application_id,
                hr_id=current_user.id,
                decision="rejected",
                decision_comments=status_update.hr_notes or "Rejected via pipeline",
                decided_at=datetime.now(timezone.utc),
            )
            db.add(decision)
    # ────────────────────────────────────────────────────────────────────

    # Atomic commit
    try:
        db.commit()
        db.refresh(application)
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail="Failed to save state transition")

    # ─── Email triggers (ONLY after successful commit) ──────────────────
    candidate_email = application.candidate_email
    job_title = application.job.title

    if result.email_type == "approved_for_interview" and raw_access_key:
        background_tasks.add_task(
            send_approved_for_interview_email, candidate_email, job_title, raw_access_key
        )
    elif result.email_type == "rejected":
        background_tasks.add_task(send_rejected_email, candidate_email, job_title, False)
    elif result.email_type == "call_for_interview":
        from app.services.email_service import send_call_for_interview_email
        background_tasks.add_task(
            send_call_for_interview_email, candidate_email, job_title
        )
    elif result.email_type == "hired":
        from app.services.email_service import send_hired_email
        background_tasks.add_task(
            send_hired_email, candidate_email, job_title, application.interview
        )
    # ────────────────────────────────────────────────────────────────────

    return {
        "id": application.id,
        "status": application.status,
        "transition": {
            "from_state": result.from_state,
            "to_state": result.to_state,
            "action": result.action,
            "email_type": result.email_type,
        }
    }


def _ensure_interview_record(application: Application, db) -> str:
    """Create or refresh Interview record + access key for an application."""
    import uuid
    raw_access_key = secrets.token_urlsafe(16)
    hashed_key = pwd_context.hash(raw_access_key)
    expiration = datetime.now(timezone.utc) + timedelta(hours=24)

    existing_interview = db.query(Interview).filter(
        Interview.application_id == application.id
    ).first()

    if not existing_interview:
        interview_stage = 'aptitude' if application.job.aptitude_enabled else 'first_level'
        unique_test_id = f"TEST-{uuid.uuid4().hex[:8].upper()}"

        new_interview = Interview(
            test_id=unique_test_id,
            application_id=application.id,
            status='not_started',
            access_key_hash=hashed_key,
            expires_at=expiration,
            is_used=False,
            interview_stage=interview_stage,
        )
        db.add(new_interview)
    elif existing_interview.status == 'not_started':
        existing_interview.access_key_hash = hashed_key
        existing_interview.expires_at = expiration
        if not existing_interview.test_id:
            existing_interview.test_id = f"TEST-{uuid.uuid4().hex[:8].upper()}"

    return raw_access_key

@router.delete("/{application_id}")
async def delete_application(
    application_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Delete an application along with associated data. HR only.
    """
    if current_user.role not in ["admin", "hr_manager", "recruiter", "hr"]:
        raise HTTPException(status_code=403, detail="Only HR or Admins can delete applications")

    app = db.query(Application).filter(Application.id == application_id).first()
    if not app:
        raise HTTPException(status_code=404, detail="Application not found")

    try:
        db.delete(app)
        db.commit()
    except Exception as e:
        db.rollback()
        print(f"Error deleting application {application_id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete application. It might have complex dependencies.")
    return {"message": "Application deleted successfully"}
@router.put("/{application_id}/notes", response_model=ApplicationResponse)
async def update_hr_notes(
    application_id: int,
    notes_update: ApplicationNotesUpdate,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Update HR notes for an application"""
    application = db.query(Application).filter(Application.id == application_id).first()
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
        
    application.hr_notes = notes_update.hr_notes
    db.commit()
    db.refresh(application)
    return application
